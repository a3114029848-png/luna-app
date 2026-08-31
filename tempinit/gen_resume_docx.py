# -*- coding: utf-8 -*-
"""生成「陆华东-标准求职简历.docx」"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------- 常量 ----------
FONT_CN = "微软雅黑"
FONT_EN = "Calibri"
PRIMARY = RGBColor(0x1F, 0x3B, 0x73)   # 深蓝
GRAY = RGBColor(0x59, 0x59, 0x59)
DARK = RGBColor(0x26, 0x26, 0x26)

doc = Document()

# ---------- 页面 ----------
sec = doc.sections[0]
sec.page_height = Cm(29.7)
sec.page_width = Cm(21.0)
sec.top_margin = Cm(1.4)
sec.bottom_margin = Cm(1.2)
sec.left_margin = Cm(1.8)
sec.right_margin = Cm(1.8)


def set_font(run, size, bold=False, color=DARK, cn=FONT_CN, en=FONT_EN):
    run.font.name = en
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), cn)


def add_para(text="", size=10, bold=False, color=DARK, align=None,
             space_before=0, space_after=2, line=1.15):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = line
    if align is not None:
        pf.alignment = align
    if text:
        r = p.add_run(text)
        set_font(r, size, bold, color)
    return p


def add_section_bar(text):
    """章节标题 + 底部横线"""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(8)
    pf.space_after = Pt(3)
    r = p.add_run(text)
    set_font(r, 12, True, PRIMARY)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), "1F3B73")
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_headline(left, right, size=10.5):
    """一行：左侧加粗标题 + 右侧时间（右对齐制表位）"""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(4)
    pf.space_after = Pt(1)
    pf.line_spacing = 1.1
    # 右对齐制表位（页面宽度）
    pf.tab_stops.add_tab_stop(Cm(17.4), WD_TAB_ALIGNMENT.RIGHT)
    r1 = p.add_run(left)
    set_font(r1, size, True, DARK)
    p.add_run("\t")
    r2 = p.add_run(right)
    set_font(r2, size, False, GRAY)
    return p


def add_bullet(text, size=10, indent=0.45):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent = Cm(indent)
    pf.first_line_indent = Cm(-0.3)
    pf.space_after = Pt(1.5)
    pf.line_spacing = 1.15
    r = p.add_run("• ")
    set_font(r, size, False, PRIMARY)
    r2 = p.add_run(text)
    set_font(r2, size, False, DARK)
    return p


# ---------- 头部 ----------
p = add_para("陆华东", 22, True, PRIMARY, WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
p = add_para("求职意向：AI 产品经理　|　22 岁　|　北京市海淀区", 10.5, False, DARK,
             WD_ALIGN_PARAGRAPH.CENTER, space_after=1)
p = add_para("131-6113-7900　|　a3114029848@163.com　|　GitHub：github.com/a3114029848-png/luna-app",
             9.5, False, GRAY, WD_ALIGN_PARAGRAPH.CENTER, space_after=4)

# ---------- 教育背景 ----------
add_section_bar("教育背景")
add_headline("北京理工大学　管理科学与工程（硕士在读）", "2025.09 – 2027.06")
add_headline("北京科技大学　信息管理与信息系统（本科）", "2021.09 – 2025.09")

# ---------- 实习经历 ----------
add_section_bar("实习经历")
add_headline("北京清研灵智科技有限公司　产品经理（AI 问卷方向）", "2024.11 – 2025.03")
add_bullet("搭建问卷字段级数据底座：清洗 3,000 条问卷数据，统一字段名/类型/编码与题型约束，"
           "产出 16 张数据字段表/导入表，将问卷隐性规范显性化为 AI 可用的结构化 Schema，为 AI 生成准确率提供数据前提")
add_bullet("设计 AI 问卷生成工作流：梳理「需求输入 → AI 草拟 → 字段校验 → 人工审校 → 导出上线」全流程，"
           "制定题型/选项/跳转生成规则与 Prompt 约束，搭建基于样例的生成质量校验，验证字段完整率与逻辑自洽性，"
           "打通 AI 生成与问卷编辑器衔接")
add_headline("国家卫生健康委卫生发展研究中心　实习助理研究员（AI 技术应用方向）", "2023.04 – 2023.10")
add_bullet("调研医疗大模型在医疗场景的应用侧重点与局限，梳理医疗 AI 评测维度"
           "（知识问答、复杂语言理解、诊疗推荐、文书生成、多轮对话、多模态交互）")
add_bullet("访谈政府、医院、企业三方角色，梳理不同组织的医疗 AI 应用需求与倾向差异")

# ---------- 项目经历 ----------
add_section_bar("项目经历")
add_headline("Luna 经期健康管理 AI 产品（独立项目）｜GitHub：github.com/a3114029848-png/luna-app",
             "2026.07 – 2026.08")
add_bullet("面向关注科学记录与就医衔接的女性用户，从 0 打造「纯本地隐私 + FIGO 医学规则 + 可溯源 AI」的经期健康 App")
add_bullet("需求与洞察：主导需求分析五步法、竞品五步法，基于公开用户原话识别「就医衔接」差异化机会，确立北极星指标「周连续记录率」")
add_bullet("AI 产品设计：设计「本地工具层 + 医学知识库 + 云端 LLM」三级架构——数据类问题走本地 FIGO 规则引擎（零成本、可溯源），"
           "科普走带来源的知识库，危险/诊断类问题 100% 兜底不交 LLM")
add_bullet("评测与合规：建立 20+ 条黄金评测集（工具命中率 ≥90%、幻觉率 ≤5%、危险兜底 =100%），"
           "设计 👍/👎 反馈闭环驱动坏样本回流，涉医回答强制免责")
add_bullet("工程落地：React Native 实现意图识别 + 4 个 Agent 工具 + 流式对话（XHR 适配 RN），真机跑通全链路；"
           "云端成本控制在约 $20/万用户/月")
add_bullet("成果：形成「需求 → 设计 → 评测 → 合规 → 数据飞轮」完整 AI 产品闭环，方法论沉淀为可复用文档与 VS Code skill")
add_headline("AI 多智能体内容生成项目（BookVoice）", "2025.07 – 2025.08")
add_bullet("基于 Coze 规划并搭建「输入解析-书籍推荐-金句生成-图片生成-素材合成」端到端自动化工作流，"
           "多次 Prompt Engineering 迭代优化生成质量与稳定性")
add_bullet("将单条视频制作时长由约 3 小时缩短至 10 分钟内，零代码跑通视频自动生成的产品化闭环")
add_headline("个人消费信贷风险预测项目", "2025.02 – 2025.05")
add_bullet("融合 Logit、SVM、随机森林模型构建借款人信用画像，基于 Kaggle 数据集（30 万行 × 122 特征）"
           "完成特征工程与模型调参，提升信贷评估准确度")

# ---------- 技能与荣誉 ----------
add_section_bar("技能与荣誉")
add_bullet("技能：熟练使用 AI 工具与工作流编排（Coze、Prompt Engineering、Agent）；数据清洗与建模（Python）；"
           "React Native 基础开发；英语 CET-4（541）/ CET-6（471）/ 雅思 6.0")
add_bullet("荣誉：北京市统计建模大赛二等奖（2023.12）")

doc.save(r"d:\Luna\陆华东-标准求职简历.docx")
print("已生成：d:\\Luna\\陆华东-标准求职简历.docx")
